"""文件文本抽取工具

支持从 PDF、DOCX、TXT 文件中提取文本内容
"""
import io
from typing import Union
import pdfplumber
from docx import Document


def detect_file_type(file_bytes: bytes, filename: str = None) -> str:
    """
    检测文件类型
    
    Args:
        file_bytes: 文件字节内容
        filename: 文件名（可选）
        
    Returns:
        文件类型 ('pdf', 'docx', 'txt')
        
    Raises:
        ValueError: 不支持的文件类型
    """
    # 首先尝试从文件名检测
    if filename:
        filename_lower = filename.lower()
        if filename_lower.endswith('.pdf'):
            return 'pdf'
        elif filename_lower.endswith(('.docx', '.doc')):
            return 'docx'
        elif filename_lower.endswith('.txt'):
            return 'txt'
    
    # 通过文件头（magic number）检测
    if len(file_bytes) < 4:
        raise ValueError("文件太小，无法检测类型")
    
    # PDF: %PDF
    if file_bytes[:4] == b'%PDF':
        return 'pdf'
    
    # DOCX: PK\x03\x04 (ZIP格式，DOCX是ZIP压缩的XML)
    if file_bytes[:2] == b'PK':
        # 检查是否是DOCX（ZIP格式，但包含特定的内部结构）
        # DOCX文件是ZIP格式，但我们可以通过尝试解析来判断
        try:
            import zipfile
            zip_file = zipfile.ZipFile(io.BytesIO(file_bytes))
            # 检查是否包含word/document.xml（DOCX的特征）
            if 'word/document.xml' in zip_file.namelist():
                return 'docx'
        except:
            pass
    
    # TXT: 尝试作为文本解码
    try:
        # 尝试UTF-8解码
        file_bytes[:1000].decode('utf-8')
        return 'txt'
    except:
        pass
    
    # 如果都检测不到，抛出错误
    raise ValueError(f"不支持的文件类型: {filename or '未知'}")


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    从 PDF 文件中提取文本
    
    Args:
        file_bytes: PDF 文件的字节内容
        
    Returns:
        提取的文本内容
        
    Raises:
        Exception: 提取失败时抛出异常
    """
    try:
        text_parts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        return "\n\n".join(text_parts)
    except Exception as e:
        raise Exception(f"PDF 文本提取失败: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    从 DOCX 文件中提取文本
    
    Args:
        file_bytes: DOCX 文件的字节内容
        
    Returns:
        提取的文本内容
        
    Raises:
        Exception: 提取失败时抛出异常
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        return "\n".join(text_parts)
    except Exception as e:
        raise Exception(f"DOCX 文本提取失败: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """
    从 TXT 文件中提取文本
    
    Args:
        file_bytes: TXT 文件的字节内容
        
    Returns:
        提取的文本内容
        
    Raises:
        Exception: 提取失败时抛出异常
    """
    try:
        # 尝试多种编码
        encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
        
        for encoding in encodings:
            try:
                text = file_bytes.decode(encoding)
                return text
            except UnicodeDecodeError:
                continue
        
        # 如果所有编码都失败，使用 utf-8 并忽略错误
        return file_bytes.decode('utf-8', errors='ignore')
    except Exception as e:
        raise Exception(f"TXT 文本提取失败: {str(e)}")


def extract_text(file_type: str = None, file_bytes: bytes = None, filename: str = None) -> str:
    """
    统一的文本提取入口
    
    Args:
        file_type: 文件类型 ('pdf', 'docx', 'txt')，如果为None则自动检测
        file_bytes: 文件的字节内容
        filename: 文件名（用于自动检测类型）
        
    Returns:
        提取的文本内容
        
    Raises:
        ValueError: 不支持的文件类型
        Exception: 提取失败时抛出异常
    """
    # 如果没有指定文件类型，自动检测
    if not file_type:
        if not file_bytes:
            raise ValueError("需要提供文件内容或文件类型")
        file_type = detect_file_type(file_bytes, filename)
    
    file_type = file_type.lower().strip()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(file_bytes)
    elif file_type == 'txt':
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"不支持的文件类型: {file_type}")
